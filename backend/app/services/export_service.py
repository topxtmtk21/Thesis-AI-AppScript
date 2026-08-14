import os
from docx import Document
from openpyxl import Workbook
import tempfile

class DocumentExporter:
    @staticmethod
    def export_md(data_list):
        md_content = "# Danh sách Tài liệu & Phân tích\n\n"
        for idx, doc in enumerate(data_list):
            md_content += f"## {idx+1}. {doc.get('title', 'Untitled')}\n"
            md_content += f"- **Tác giả & Năm**: {doc.get('authors', '')} ({doc.get('year', '')})\n"
            md_content += f"- **Tạp chí**: {doc.get('journal', '')}\n"
            md_content += f"- **Lý thuyết**: {doc.get('theory', '')}\n"
            md_content += f"- **Phương pháp**: {doc.get('methodology', '')}\n"
            
            findings = doc.get("detailedFindings", [])
            if findings:
                md_content += "- **Các Nội dung cốt lõi**:\n"
                for f in findings:
                    md_content += f"  - **[{f.get('location', '')}]** {f.get('content', '')}\n"
            else:
                md_content += f"- **Kết quả**: {doc.get('keyFindings', '')}\n"
                
            md_content += f"- **Khoảng trống**: {doc.get('researchGap', '')}\n"
            md_content += "---\n"
        return md_content

    @staticmethod
    def export_docx(data_list):
        temp_file = tempfile.NamedTemporaryFile(delete=False, suffix=".docx")
        doc = Document()
        doc.add_heading('Danh sách Tài liệu & Phân tích', 0)
        
        for idx, item in enumerate(data_list):
            doc.add_heading(f"{idx+1}. {item.get('title', 'Untitled')}", level=1)
            doc.add_paragraph(f"Tác giả & Năm: {item.get('authors', '')} ({item.get('year', '')})")
            doc.add_paragraph(f"Tạp chí: {item.get('journal', '')}")
            doc.add_paragraph(f"Lý thuyết: {item.get('theory', '')}")
            doc.add_paragraph(f"Phương pháp: {item.get('methodology', '')}")
            
            findings = item.get("detailedFindings", [])
            if findings:
                doc.add_paragraph("Các Nội dung cốt lõi:")
                for f in findings:
                    doc.add_paragraph(f"[{f.get('location', '')}] {f.get('content', '')}", style='List Bullet')
            else:
                doc.add_paragraph(f"Kết quả: {item.get('keyFindings', '')}")
            doc.add_paragraph(f"Tài liệu tham khảo (References):")
            
            refs = item.get("references", [])
            if refs:
                for r in refs:
                    doc.add_paragraph(r, style='List Bullet')
            else:
                doc.add_paragraph("Không có")
                
            doc.add_page_break()
            
        doc.save(temp_file.name)
        return temp_file.name

    @staticmethod
    def export_xlsx(data_list):
        temp_file = tempfile.NamedTemporaryFile(delete=False, suffix=".xlsx")
        wb = Workbook()
        ws = wb.active
        ws.title = "References Data"
        
        # Headers
        headers = ["STT", "Tựa đề", "Tác giả", "Năm", "Tạp chí", "Lý thuyết", "Phương pháp", "Nội dung cốt lõi (Chi tiết)", "References"]
        ws.append(headers)
        
        for idx, item in enumerate(data_list):
            refs = item.get("references", [])
            refs_str = "\n".join(refs) if refs else "Không có"
            
            findings = item.get("detailedFindings", [])
            if findings:
                findings_str = "\n".join([f"[{f.get('location', '')}] {f.get('content', '')}" for f in findings])
            else:
                findings_str = item.get("keyFindings", "")
            
            row = [
                idx + 1,
                item.get('title', ''),
                item.get('authors', ''),
                item.get('year', ''),
                item.get('journal', ''),
                item.get('theory', ''),
                item.get('methodology', ''),
                findings_str,
                refs_str
            ]
            ws.append(row)
            
        wb.save(temp_file.name)
        return temp_file.name
